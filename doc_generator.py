"""
Word Document Generator
生成小麦抗病·植物免疫论文日报
"""

from __future__ import annotations

import logging
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from analyzer import PaperAnalysis
from config import (
    DOCUMENT_TITLE,
    DOCUMENT_AUTHOR,
    DOCUMENT_FONT,
    OUTPUT_DIR,
    SUMMARY_SECTIONS,
)

logger = logging.getLogger(__name__)

# ── 配色方案（小麦金 + 深绿）──────────────────────────────────────────────
C_WHEAT   = RGBColor(0xC8, 0x96, 0x20)   # 小麦金
C_GREEN   = RGBColor(0x1A, 0x6B, 0x2F)   # 深绿
C_DARK    = RGBColor(0x22, 0x22, 0x22)   # 正文深灰
C_MID     = RGBColor(0x55, 0x55, 0x55)   # 次要信息
C_LIGHT   = RGBColor(0x99, 0x99, 0x99)   # 更浅
C_RED     = RGBColor(0xCC, 0x33, 0x00)   # 错误/警告
C_SCORE_HI = RGBColor(0xC8, 0x60, 0x00)  # 高分标记（橙）
C_SCORE_MID = RGBColor(0x1A, 0x6B, 0x2F) # 中分标记（绿）


def _set_font(run, size_pt: float, bold=False, color: RGBColor = None, font=None):
    run.font.name = font or DOCUMENT_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 中文字体
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font or DOCUMENT_FONT)
    rPr.insert(0, rFonts)


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph(text)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p


class DocumentGenerator:

    def __init__(self, output_dir: Optional[Path] = None):
        self.output_dir = output_dir or OUTPUT_DIR

    def generate(self, analyses: list[PaperAnalysis], fetch_date: Optional[date] = None) -> Path:
        fetch_date = fetch_date or date.today()
        doc = Document()
        self._page_setup(doc)
        self._setup_styles(doc)

        # 封面
        self._cover(doc, analyses, fetch_date)
        doc.add_page_break()

        # 论文简报（每篇一节，不强制分页，保持连贯阅读）
        valid   = [a for a in analyses if a.is_valid()]
        invalid = [a for a in analyses if not a.is_valid()]

        for i, analysis in enumerate(valid):
            self._paper_block(doc, analysis, i + 1, len(valid))

        if invalid:
            doc.add_page_break()
            self._failed_appendix(doc, invalid)

        # 页脚说明
        self._footer_note(doc, len(valid), len(invalid), fetch_date)

        # 保存
        self.output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"小麦抗病论文日报_{fetch_date.strftime('%Y%m%d')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        logger.info("Document saved: %s", filepath)
        return filepath

    # ── 页面与样式 ──────────────────────────────────────────────────────────

    def _page_setup(self, doc):
        sec = doc.sections[0]
        sec.page_width   = Cm(21.0)
        sec.page_height  = Cm(29.7)
        sec.top_margin   = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin  = Cm(2.8)
        sec.right_margin = Cm(2.8)

    def _setup_styles(self, doc):
        normal = doc.styles['Normal']
        normal.font.name = DOCUMENT_FONT
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = C_DARK
        normal.paragraph_format.space_after  = Pt(4)
        normal.paragraph_format.line_spacing = 1.4
        rPr = normal.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), DOCUMENT_FONT)
        rPr.insert(0, rFonts)

        for lvl, size, bold in [(1, 16, True), (2, 13, True), (3, 11, True)]:
            st = doc.styles[f'Heading {lvl}']
            st.font.name  = DOCUMENT_FONT
            st.font.size  = Pt(size)
            st.font.bold  = bold
            st.font.color.rgb = C_GREEN
            rPr = st.element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), DOCUMENT_FONT)
            rPr.insert(0, rFonts)

    # ── 封面 ────────────────────────────────────────────────────────────────

    def _cover(self, doc, analyses, fetch_date: date):
        for _ in range(5):
            doc.add_paragraph()

        # 主标题
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        r = p.add_run("🌾 小麦抗病·植物免疫")
        _set_font(r, 26, bold=True, color=C_WHEAT)

        p2 = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        r2 = p2.add_run("论文日报")
        _set_font(r2, 22, bold=True, color=C_GREEN)

        # 日期
        p3 = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
        r3 = p3.add_run(fetch_date.strftime("%Y年%m月%d日"))
        _set_font(r3, 13, color=C_MID)

        # 分隔线
        sep = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
        sep.add_run("─" * 38).font.color.rgb = C_LIGHT

        # 统计
        valid = [a for a in analyses if a.is_valid()]
        hi    = [a for a in valid if a.relevance_score >= 8]
        mid   = [a for a in valid if 6 <= a.relevance_score < 8]

        stats = [
            f"📄 本期收录论文：{len(valid)} 篇",
            f"🎯 高度相关（小麦抗病）：{len(hi)} 篇",
            f"🌿 相关（植物抗病/免疫）：{len(mid)} 篇",
            f"📅 文献检索范围：近 3 天",
            f"📰 数据来源：bioRxiv · PubMed · Nature Plants · Plant Cell · New Phytologist 等",
            f"🤖 分析引擎：DeepSeek AI",
            f"⏰ 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        ]
        for s in stats:
            p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            r = p.add_run(s)
            _set_font(r, 10.5, color=C_DARK)

    # ── 单篇论文区块 ────────────────────────────────────────────────────────

    def _paper_block(self, doc, analysis: PaperAnalysis, idx: int, total: int):
        doc.add_paragraph()  # 间隔

        paper = analysis.paper
        score = analysis.relevance_score

        # ── 标题行 ──────────────────────────────────────────────────────────
        title_p = doc.add_heading(f"{idx}. {paper.title}", level=2)

        # ── 相关性徽章 ───────────────────────────────────────────────────────
        badge_p = _para(doc, space_after=6)
        score_color = C_SCORE_HI if score >= 8 else C_SCORE_MID
        badge_r = badge_p.add_run(f"  ★ 相关性评分：{score}/10  ")
        _set_font(badge_r, 10, bold=True, color=score_color)
        if analysis.relevance_note:
            note_r = badge_p.add_run(f"— {analysis.relevance_note.strip()}")
            _set_font(note_r, 9.5, color=C_MID)

        # ── 基本信息一行 ─────────────────────────────────────────────────────
        journal  = getattr(paper, 'source', '') or getattr(paper, 'journal', '') or 'N/A'
        pub_date = paper.published_date.isoformat() if paper.published_date else 'N/A'
        doi      = getattr(paper, 'doi', '') or 'N/A'
        url      = getattr(paper, 'url', '') or ''
        authors  = getattr(paper, 'author_str', '') or ''

        meta_items = [
            ("📰 期刊", journal),
            ("📅 日期", pub_date),
            ("👤 作者", authors[:120] + ("…" if len(authors) > 120 else "")),
            ("🔗 DOI",  doi),
        ]
        if url and url != 'N/A':
            meta_items.append(("🌐 链接", url))

        for label, value in meta_items:
            if not value or value == 'N/A':
                continue
            mp = _para(doc, space_before=0, space_after=2)
            mp.paragraph_format.left_indent = Cm(0.3)
            lr = mp.add_run(f"{label}：")
            _set_font(lr, 9.5, bold=True, color=C_GREEN)
            vr = mp.add_run(value)
            _set_font(vr, 9.5, color=C_MID)

        doc.add_paragraph()

        # ── 各分析节 ─────────────────────────────────────────────────────────
        section_map = {
            "publication_info": ("📋 发表信息",   analysis.publication_info),
            "affiliation":      ("🏛  研究单位",   analysis.affiliation),
            "background":       ("🔬 研究背景",   analysis.background),
            "methods":          ("🧪 研究方法",   analysis.methods),
            "results":          ("📊 实验结果",   analysis.results),
            "discussion":       ("💬 讨论",       analysis.discussion),
            "innovations":      ("⭐ 创新点",     analysis.innovations),
            "relevance_note":   ("🌾 与小麦抗病的关联", analysis.relevance_note),
        }

        for key, label in SUMMARY_SECTIONS:
            _, content = section_map.get(key, (None, ""))
            if not content or not content.strip():
                continue

            # 小节标题
            sh = doc.add_heading(label, level=3)

            # 内容
            cp = _para(doc, space_before=0, space_after=6)
            cp.paragraph_format.left_indent = Cm(0.5)
            cr = cp.add_run(content.strip())
            _set_font(cr, 10.5, color=C_DARK)

        # ── 分隔线 ───────────────────────────────────────────────────────────
        sep = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
        sep.add_run("· · · · ·").font.color.rgb = C_LIGHT

    # ── 失败附录 ────────────────────────────────────────────────────────────

    def _failed_appendix(self, doc, failed: list[PaperAnalysis]):
        doc.add_heading("📋 附录：未能完整分析的论文", level=1)
        note = doc.add_paragraph(f"以下 {len(failed)} 篇论文相关性较低或分析出错，仅列出基本信息供参考。")
        for run in note.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = C_MID

        for i, a in enumerate(failed, 1):
            p = _para(doc, space_after=3)
            p.paragraph_format.left_indent = Cm(0.3)
            ir = p.add_run(f"{i}. ")
            _set_font(ir, 10, bold=True)
            tr = p.add_run(f"{a.paper.title}  ")
            _set_font(tr, 10)
            sr = p.add_run(f"[{getattr(a.paper, 'source', 'N/A')} | {a.paper.published_date or 'N/A'} | 相关性: {a.relevance_score}/10]")
            _set_font(sr, 9, color=C_LIGHT)

            err = getattr(a, 'analysis_error', '')
            if err:
                ep = _para(doc, space_after=2)
                ep.paragraph_format.left_indent = Cm(1.0)
                er = ep.add_run(f"错误: {err}")
                _set_font(er, 9, color=C_RED)

    # ── 页脚说明 ────────────────────────────────────────────────────────────

    def _footer_note(self, doc, success: int, failed: int, fetch_date: date):
        doc.add_paragraph()
        doc.add_heading("📝 说明", level=2)
        notes = [
            "本简报由 Paper Digest Bot 自动检索、AI 分析生成，每日早 8:00 更新。",
            "数据来源：bioRxiv · PubMed · Nature Plants · The Plant Cell · New Phytologist · Nature · Cell 等。",
            "筛选标准：优先收录小麦抗病（条锈、叶锈、白粉、赤霉等）及植物免疫机制相关研究性论文。",
            "分析引擎：DeepSeek AI。内容仅供参考，详细内容请查阅原文。",
            f"本期统计：成功分析 {success} 篇，低相关/失败 {failed} 篇。",
        ]
        for nt in notes:
            np = _para(doc, space_after=2)
            nr = np.add_run(nt)
            _set_font(nr, 9, color=C_LIGHT)
